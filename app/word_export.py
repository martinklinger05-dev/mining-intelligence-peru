from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH


def ensure_docs_folder() -> None:
    Path("docs").mkdir(parents=True, exist_ok=True)


def export_accident_report_to_docx(report: dict[str, Any]) -> str:
    """
    report expected keys:
      - timestamp
      - description
      - acto_subestandar
      - condicion_subestandar
      - consequence (optional)
      - causa_inmediata (optional)
      - cinco_porques (optional list[str])
      - normative_support
      - actions (optional list[dict])
      - full_text (optional)  # if you only have the big report text
    """
    ensure_docs_folder()
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_path = Path("docs") / f"informe_investigacion_{ts}.docx"

    doc = Document()

    # Title
    title = doc.add_heading("INFORME DE INVESTIGACIÓN DE ACCIDENTE", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Header metadata (mine style)
    meta = report.get("meta", {})

    doc.add_paragraph(f"Área / Labor: {meta.get('area', '__________________________')}")
    doc.add_paragraph(f"Equipo: {meta.get('equipo', '__________________________')}")
    doc.add_paragraph(f"Turno: {meta.get('turno', '__________________________')}")
    doc.add_paragraph(f"Fecha/Hora del evento: {meta.get('fecha_evento', '__________________________')}")
    doc.add_paragraph(f"Lesiones: {meta.get('lesiones', '__________________________')}")
    doc.add_paragraph(f"Tipo de daño: {meta.get('danio', '__________________________')}")
    doc.add_paragraph("Supervisor Responsable: __________________________")
    doc.add_paragraph("Código de Informe: __________________________")

    # Sections (structured if available)
    doc.add_heading("1. RESUMEN DEL EVENTO", level=2)
    doc.add_paragraph(report.get("description", ""))

    doc.add_heading("2. HALLAZGO CLAVE", level=2)
    acto = report.get("acto_subestandar", "N/D")
    condicion = report.get("condicion_subestandar", "N/D")
    doc.add_paragraph(f"El evento se asocia a {acto.lower()} en presencia de {condicion.lower()}.")

    doc.add_heading("3. CLASIFICACIÓN", level=2)
    doc.add_paragraph(f"Acto subestándar: {acto}")
    doc.add_paragraph(f"Condición subestándar: {condicion}")
    if "consequence" in report and report["consequence"]:
        doc.add_paragraph(f"Consecuencia: {report['consequence']}")

    doc.add_heading("4. CAUSA INMEDIATA", level=2)
    doc.add_paragraph(report.get("causa_inmediata", "N/D"))

    doc.add_heading("5. CAUSA RAÍZ (5 POR QUÉS)", level=2)
    cinco = report.get("cinco_porques", [])
    if isinstance(cinco, list) and cinco:
        for line in cinco:
            doc.add_paragraph(line)
    else:
        doc.add_paragraph("N/D")

    doc.add_heading("6. SUSTENTO NORMATIVO (DS 024-2016-EM)", level=2)
    doc.add_paragraph(report.get("normative_support", "N/D"))

    doc.add_heading("7. ACCIONES CORRECTIVAS / PREVENTIVAS", level=2)

    actions = report.get("actions", [])
    if isinstance(actions, list) and actions:
        table = doc.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        hdr[0].text = "N°"
        hdr[1].text = "Acción"
        hdr[2].text = "Responsable"
        hdr[3].text = "Plazo"
        table.style = "Table Grid"

        for i, a in enumerate(actions, 1):
            row = table.add_row().cells
            row[0].text = str(i)
            row[1].text = a.get("accion", "")
            row[2].text = a.get("responsable", "")
            row[3].text = a.get("plazo", "")
    else:
        # fallback: if you only have a full text block
        full_text = report.get("full_text", "")
        doc.add_paragraph(full_text if full_text else "N/D")

    doc.add_paragraph("")
    doc.add_paragraph("_______________________________")
    doc.add_paragraph("Firma Supervisor Responsable")

    doc.save(str(out_path))
    return str(out_path)